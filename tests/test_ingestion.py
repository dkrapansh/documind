import io

from app.models.chunk import Chunk
from app.models.document import Document

def _fake_embed_text(text: str) -> list[float]:
    """Stand-in for the real OpenAI call — instant, free, deterministic.
    Ingestion only needs *a* vector of the right shape to prove the
    pipeline wires together; it doesn't need a real embedding to test
    chunking, DB writes, or status transitions."""
    return [0.1] * 1536

def _auth_headers(client) -> dict:
    response = client.post("/auth/keys", json={"tenant_name": "acme"})
    raw_key = response.json()["api_key"]
    return {"X-API-Key": raw_key}

def test_upload_ingests_document_into_chunks_with_embeddings(client, db_session, monkeypatch):
    monkeypatch.setattr("app.services.ingestion.embed_text", _fake_embed_text)
    headers = _auth_headers(client)

    file_content = b"Hello world. This is a small test document for ingestion."
    response = client.post(
        "/documents",
        headers=headers,
        files={"file": ("ingestion_test.txt", io.BytesIO(file_content), "text/plain")},
    )
    assert response.status_code == 200
    document_id = response.json()["id"]

    status_response = client.get(f"/documents/{document_id}", headers=headers)
    body = status_response.json()
    assert body["status"] == "ready"
    assert body["chunk_count"] == 1  # short text stays under one chunk

    chunks = db_session.query(Chunk).filter(Chunk.document_id == document_id).all()
    assert len(chunks) == 1
    assert chunks[0].text.strip() == file_content.decode("utf-8").strip()
    assert len(chunks[0].embedding) == 1536

def test_upload_docx_extracts_paragraph_and_table_text(client, db_session, monkeypatch):
    """.docx is one of the 3 supported extensions (services/text_extraction.py)
    but until now had no test exercising it through the real pipeline - table
    cell text specifically, which is appended separately from paragraph text
    and previously crashed process_document (a `cell.text.sstrip()` typo,
    silently swallowed by ingestion.py's except-Exception into a "failed"
    status rather than a visible traceback)."""
    monkeypatch.setattr("app.services.ingestion.embed_text", _fake_embed_text)
    headers = _auth_headers(client)

    from docx import Document as DocxDocument

    docx_document = DocxDocument()
    docx_document.add_paragraph("Refund Policy: refunds within 30 days.")
    table = docx_document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Region"
    table.rows[0].cells[1].text = "EU"
    buffer = io.BytesIO()
    docx_document.save(buffer)
    buffer.seek(0)

    response = client.post(
        "/documents",
        headers=headers,
        files={
            "file": (
                "policy.docx",
                buffer,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    document_id = response.json()["id"]

    status = client.get(f"/documents/{document_id}", headers=headers).json()
    assert status["status"] == "ready"

    chunks = db_session.query(Chunk).filter(Chunk.document_id == document_id).all()
    combined_text = " ".join(chunk.text for chunk in chunks)
    assert "Refund Policy" in combined_text
    assert "Region" in combined_text and "EU" in combined_text

def test_reupload_same_file_is_a_noop(client, db_session, monkeypatch):
    monkeypatch.setattr("app.services.ingestion.embed_text", _fake_embed_text)
    headers = _auth_headers(client)

    file_content = b"Duplicate upload content for idempotency test."

    first = client.post(
        "/documents",
        headers=headers,
        files={"file": ("dup_test.txt", io.BytesIO(file_content), "text/plain")},
    )
    second = client.post(
        "/documents",
        headers=headers,
        files={"file": ("dup_test.txt", io.BytesIO(file_content), "text/plain")},
    )

    assert first.json()["id"] == second.json()["id"]

    documents = db_session.query(Document).all()
    assert len(documents) == 1

    chunks = db_session.query(Chunk).filter(Chunk.document_id == first.json()["id"]).all()
    assert len(chunks) == 1