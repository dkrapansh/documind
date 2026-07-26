import io 
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

from app.core.exceptions import UnsupportedFileTypeException

SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}

def validate_extension(filename: str) -> None:
    """Reject unsupported file types at upload time, before saving to
    disk or scheduling a background job that would only fail later
    with no feedback to the client."""
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileTypeException(extension)
    

def extract_text(file_bytes: bytes, filename: str) -> str:
    """Turn a saved file's raw bytes into plain text, based on extension.
    .txt  -> decode as UTF-8 directly.
    .pdf  -> pypdf's text layer. Scanned/image-only PDFs extract as
             empty string; OCR is a deliberate out-of-scope limitation.
    .docx -> paragraph text plus table cell text via python-docx,
             tables appended after paragraphs rather than at their
             true position, fine since chunking cares about content,
             not reading order.
    """
    extension = Path(filename).suffix.lower()

    if extension == ".txt":
        return file_bytes.decode("utf-8")
    
    if extension == ".pdf":
        reader = PdfReader(io.BytesIO(file_bytes))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n\n".join(pages)
    
    if extension == ".docx":
        doc = DocxDocument(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        table_cells = [ 
            cell.text
            for table in doc.tables
            for row in table.rows
            for cell in row.cells
            if cell.text.strip()
        ]
        return "\n\n".join(paragraphs + table_cells)
    
    raise UnsupportedFileTypeException(extension)